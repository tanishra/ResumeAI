import io
from typing import Tuple
import pdfplumber
from docx import Document

def extract_text_from_pdf(file_bytes: bytes) -> str:
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            parts.append(txt)
    return "\n".join(parts)

def extract_text_from_docx(file_bytes: bytes) -> str:
    f = io.BytesIO(file_bytes)
    doc = Document(f)
    parts = []
    
    # Extract from paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                parts.append(" | ".join(row_text))
                
    return "\n".join(parts)

def detect_and_extract(filename: str, file_bytes: bytes) -> Tuple[str, str]:
    """Return (ext, text). ext in {pdf, docx, txt}."""
    low = filename.lower()
    if low.endswith(".pdf"):
        return "pdf", extract_text_from_pdf(file_bytes)
    if low.endswith(".docx"):
        return "docx", extract_text_from_docx(file_bytes)
    # basic text fallback
    try:
        return "txt", file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        return "bin", ""
