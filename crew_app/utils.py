import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# --- DOCX GENERATOR ---
def txt_to_docx_bytes(text: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = [line.strip() for line in text.splitlines()]
    if not lines: return b""

    # Name Handling
    name_line = ""
    for i, line in enumerate(lines):
        if line:
            name_line = line
            lines = lines[i+1:]
            break
    if name_line:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name_line)
        run.bold = True
        run.font.size = Pt(22)

    # Content Handling
    for line in lines:
        if not line:
            doc.add_paragraph()
            continue
        is_header = line.isupper() and len(line) < 40 and not any(c in line for c in ['|', '@', ':', '/'])
        if is_header:
            p = doc.add_paragraph()
            p.space_before = Pt(12)
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(13)
            continue
        if line.startswith(('-', '•', '*')):
            p = doc.add_paragraph(style='List Bullet')
            clean_line = re.sub(r'^[-•*]\s*', '', line)
            p.add_run(clean_line)
        else:
            doc.add_paragraph(line)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# --- PDF TEMPLATE ENGINE ---
class ResumeTemplate(FPDF):
    def __init__(self):
        super().__init__()
        self.l_margin_val = 20
        self.r_margin_val = 20
        self.set_margins(self.l_margin_val, 20, self.r_margin_val)
        self.set_auto_page_break(auto=True, margin=20)
        self.indigo = (79, 70, 229)
        self.slate = (15, 23, 42)
        self.muted = (100, 116, 139)
        self.content_w = 210 - self.l_margin_val - self.r_margin_val

    def draw_section_header(self, title):
        self.ln(6)
        self.set_font("helvetica", "B", 11)
        self.set_text_color(*self.indigo)
        self.cell(self.content_w, 8, title.upper(), ln=True)
        # Decorative Line
        curr_x, curr_y = self.get_x(), self.get_y()
        self.set_draw_color(*self.indigo)
        self.set_line_width(0.3)
        self.line(curr_x, curr_y, curr_x + self.content_w, curr_y)
        self.ln(3)
        self.set_text_color(*self.slate)

    def draw_name(self, name):
        self.set_font("helvetica", "B", 26)
        self.set_text_color(*self.slate)
        self.cell(self.content_w, 15, name, ln=True, align="C")
        self.ln(2)

    def draw_contact_info(self, info_line):
        self.set_font("helvetica", "", 9)
        self.set_text_color(*self.muted)
        self.cell(self.content_w, 5, info_line, ln=True, align="C")
        self.ln(5)

def _clean_text(text: str) -> str:
    """Sanitize text for PDF encoding."""
    replacements = {
        "\u2013": "-", "\u2014": "-", "\u2018": "'", "\u2019": "'",
        "\u201c": '"', "\u201d": '"', "\u2022": "-", "\u2026": "..."
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode("latin-1", "replace").decode("latin-1")

def txt_to_pdf_bytes(text: str) -> bytes:
    cleaned_text = _clean_text(text)
    pdf = ResumeTemplate()
    pdf.add_page()
    
    lines = [line.strip() for line in cleaned_text.splitlines()]
    if not lines: return b""

    content_start_idx = 0
    for i, line in enumerate(lines):
        if line:
            pdf.draw_name(line)
            content_start_idx = i + 1
            break
    
    contact_info = []
    for i in range(content_start_idx, min(content_start_idx + 2, len(lines))):
        if lines[i] and ("@" in lines[i] or "|" in lines[i] or any(c.isdigit() for c in lines[i])):
            contact_info.append(lines[i])
            content_start_idx = i + 1
    
    if contact_info:
        pdf.draw_contact_info("  |  ".join(contact_info))

    for i in range(content_start_idx, len(lines)):
        line = lines[i]
        if not line:
            pdf.ln(2)
            continue

        if line.isupper() and len(line) < 35:
            pdf.draw_section_header(line)
            continue

        is_bold_line = "|" in line or re.search(r'\b(19|20)\d{2}\b', line)
        if is_bold_line and not line.startswith('-'):
            pdf.set_font("helvetica", "B", 10)
            pdf.multi_cell(pdf.content_w, 6, line)
            continue

        pdf.set_font("helvetica", "", 10)
        if line.startswith('-'):
            # Safe Indentation Logic
            indent = 6
            pdf.set_x(pdf.l_margin_val + indent)
            pdf.multi_cell(pdf.content_w - indent, 6, line)
        else:
            pdf.multi_cell(pdf.content_w, 6, line)

    # FINAL BINARY PASS: Starlette/FastAPI require bytes
    try:
        # Check if output returns bytes (fpdf2) or needs manual handling (fpdf)
        raw_pdf = pdf.output()
        if isinstance(raw_pdf, (bytearray, bytes)):
            return bytes(raw_pdf)
        # If it returns a string (older versions with dest='S')
        output = pdf.output(dest='S')
        return output.encode('latin-1') if isinstance(output, str) else bytes(output)
    except Exception:
        # Extreme fallback for binary extraction
        output = pdf.output(dest='S')
        if isinstance(output, str): return output.encode('latin-1')
        return bytes(output)
