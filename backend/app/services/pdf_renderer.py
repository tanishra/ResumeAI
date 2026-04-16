from __future__ import annotations

from dataclasses import asdict, dataclass, field
from html import escape
from io import BytesIO
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


SECTION_ALIASES = {
    "summary": "summary",
    "professional summary": "summary",
    "profile": "summary",
    "education": "education",
    "academic background": "education",
    "experience": "experience",
    "work experience": "experience",
    "professional experience": "experience",
    "employment": "experience",
    "projects": "projects",
    "project experience": "projects",
    "technical skills": "skills",
    "skills": "skills",
    "core skills": "skills",
    "certifications": "certifications",
    "certification": "certifications",
    "licenses": "certifications",
    "extracurricular": "extracurricular",
    "extracurricular activities": "extracurricular",
    "activities": "extracurricular",
    "leadership": "extracurricular",
    "achievements": "extracurricular",
}

SECTION_ORDER = [
    "summary",
    "education",
    "experience",
    "projects",
    "skills",
    "certifications",
    "extracurricular",
]

DATE_RE = re.compile(
    r"(?i)\b("
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{4}"
    r"(?:\s*[-–—]\s*(?:present|current|"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?\s+\d{4}|\d{4}))?"
    r"|(?:19|20)\d{2}\s*[-–—]\s*(?:present|current|(?:19|20)\d{2})"
    r"|(?:19|20)\d{2}"
    r")\b"
)
EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(?<!\w)(?:\+?\d[\d\s().-]{7,}\d)")
URL_RE = re.compile(r"(https?://[^\s|]+|www\.[^\s|]+)")


@dataclass
class ResumeEntry:
    title: str = ""
    date: str = ""
    subtitle: str = ""
    location: str = ""
    bullets: list[str] = field(default_factory=list)


@dataclass
class ResumeDocument:
    name: str
    location: str = ""
    phone: str = ""
    email: str = ""
    links: list[tuple[str, str]] = field(default_factory=list)
    summary: list[str] = field(default_factory=list)
    education: list[ResumeEntry] = field(default_factory=list)
    experience: list[ResumeEntry] = field(default_factory=list)
    projects: list[ResumeEntry] = field(default_factory=list)
    skills: list[tuple[str, str]] = field(default_factory=list)
    certifications: list[str] = field(default_factory=list)
    extracurricular: list[str] = field(default_factory=list)


def structured_resume_from_text(final_resume_text: str) -> dict:
    return asdict(parse_resume_text(final_resume_text))


def structured_resume_from_payload(payload: dict | None, final_resume_text: str = "") -> dict:
    if payload:
        return payload
    return structured_resume_from_text(final_resume_text)


def resume_document_from_payload(payload: dict | None, final_resume_text: str = "") -> ResumeDocument:
    if not payload:
        return parse_resume_text(final_resume_text)

    return ResumeDocument(
        name=str(payload.get("name", "")).strip() or "Candidate Name",
        location=str(payload.get("location", "")).strip(),
        phone=str(payload.get("phone", "")).strip(),
        email=str(payload.get("email", "")).strip(),
        links=[(str(label), str(url)) for label, url in payload.get("links", []) if label and url],
        summary=[str(item).strip() for item in payload.get("summary", []) if str(item).strip()],
        education=_entries_from_payload(payload.get("education", [])),
        experience=_entries_from_payload(payload.get("experience", [])),
        projects=_entries_from_payload(payload.get("projects", [])),
        skills=[
            (str(label).strip(), str(value).strip())
            for label, value in payload.get("skills", [])
            if str(label).strip() and str(value).strip()
        ],
        certifications=[str(item).strip() for item in payload.get("certifications", []) if str(item).strip()],
        extracurricular=[str(item).strip() for item in payload.get("extracurricular", []) if str(item).strip()],
    )


def render_resume_pdf_bytes(final_resume_text: str = "", structured_resume: dict | None = None) -> bytes:
    parsed = resume_document_from_payload(structured_resume, final_resume_text)
    return _build_pdf(parsed)


def render_resume_docx_bytes(final_resume_text: str = "", structured_resume: dict | None = None) -> bytes:
    parsed = resume_document_from_payload(structured_resume, final_resume_text)
    return _build_docx(parsed)


def render_resume_html(final_resume_text: str = "", structured_resume: dict | None = None) -> str:
    parsed = resume_document_from_payload(structured_resume, final_resume_text)
    return _build_html(parsed)


def parse_resume_text(final_resume_text: str) -> ResumeDocument:
    lines = [_strip_markdown_artifacts(line.rstrip()) for line in final_resume_text.splitlines()]
    header_lines, section_map = _split_resume_sections(lines)

    header = _parse_header(header_lines)
    return ResumeDocument(
        name=header["name"],
        location=header["location"],
        phone=header["phone"],
        email=header["email"],
        links=header["links"],
        summary=_parse_simple_lines(section_map.get("summary", [])),
        education=_parse_standard_entries(section_map.get("education", []), kind="education"),
        experience=_parse_standard_entries(section_map.get("experience", []), kind="experience"),
        projects=_parse_standard_entries(section_map.get("projects", []), kind="projects"),
        skills=_parse_skills(section_map.get("skills", [])),
        certifications=_parse_simple_lines(section_map.get("certifications", [])),
        extracurricular=_parse_simple_lines(section_map.get("extracurricular", [])),
    )


def _entries_from_payload(items: list[dict]) -> list[ResumeEntry]:
    entries: list[ResumeEntry] = []
    for item in items:
        entries.append(
            ResumeEntry(
                title=str(item.get("title", "")).strip(),
                date=str(item.get("date", "")).strip(),
                subtitle=str(item.get("subtitle", "")).strip(),
                location=str(item.get("location", "")).strip(),
                bullets=[str(bullet).strip() for bullet in item.get("bullets", []) if str(bullet).strip()],
            )
        )
    return entries


def _build_html(parsed: ResumeDocument) -> str:
    def para(text: str) -> str:
        return f"<p>{escape(text)}</p>"

    sections: list[str] = []
    if parsed.summary:
        sections.append(f"<section><h2>Summary</h2>{para(' '.join(parsed.summary))}</section>")
    if parsed.education:
        sections.append(_entries_html("Education", parsed.education))
    if parsed.experience:
        sections.append(_entries_html("Experience", parsed.experience))
    if parsed.projects:
        sections.append(_entries_html("Projects", parsed.projects))
    if parsed.skills:
        skills = "".join(
            f"<li><strong>{escape(label)}:</strong> {escape(value)}</li>"
            for label, value in parsed.skills
        )
        sections.append(f"<section><h2>Technical Skills</h2><ul>{skills}</ul></section>")
    if parsed.certifications:
        certs = "".join(f"<li>{escape(item)}</li>" for item in parsed.certifications)
        sections.append(f"<section><h2>Certifications</h2><ul>{certs}</ul></section>")
    if parsed.extracurricular:
        items = "".join(f"<li>{escape(item)}</li>" for item in parsed.extracurricular)
        sections.append(f"<section><h2>Extracurricular</h2><ul>{items}</ul></section>")

    contacts: list[str] = []
    if parsed.phone:
        contacts.append(escape(parsed.phone))
    if parsed.email:
        contacts.append(escape(parsed.email))
    if parsed.location:
        contacts.append(escape(parsed.location))
    contacts.extend(f'<a href="{escape(url)}">{escape(label)}</a>' for label, url in parsed.links)

    return f"""
<!doctype html>
<html>
<head>
  <meta charset="utf-8" />
  <title>{escape(parsed.name)}</title>
  <style>
    body {{ font-family: Helvetica, Arial, sans-serif; color: #0f172a; margin: 40px; line-height: 1.45; }}
    h1 {{ text-align: center; margin: 0; font-size: 28px; letter-spacing: 0.08em; text-transform: uppercase; }}
    .meta {{ text-align: center; color: #475569; font-size: 11px; margin: 8px 0 20px; }}
    .meta a {{ color: #0f766e; text-decoration: none; }}
    h2 {{ font-size: 13px; margin: 18px 0 8px; text-transform: uppercase; letter-spacing: 0.12em; border-bottom: 1px solid #cbd5e1; padding-bottom: 4px; }}
    .entry-head {{ display: flex; justify-content: space-between; gap: 12px; font-weight: 700; }}
    .entry-sub {{ display: flex; justify-content: space-between; gap: 12px; color: #334155; font-style: italic; margin-bottom: 4px; }}
    ul {{ margin: 6px 0 12px 18px; padding: 0; }}
    li {{ margin: 2px 0; }}
    p {{ margin: 0 0 8px; }}
  </style>
</head>
<body>
  <h1>{escape(parsed.name)}</h1>
  <div class="meta">{' | '.join(contacts)}</div>
  {''.join(sections)}
</body>
</html>
""".strip()


def _entries_html(title: str, entries: list[ResumeEntry]) -> str:
    blocks: list[str] = [f"<section><h2>{escape(title)}</h2>"]
    for entry in entries:
        head_left = escape(entry.title)
        head_right = escape(entry.date)
        sub_left = escape(entry.subtitle)
        sub_right = escape(entry.location)
        bullets = "".join(f"<li>{escape(item)}</li>" for item in entry.bullets)
        blocks.append(
            f'<div class="entry">'
            f'<div class="entry-head"><span>{head_left}</span><span>{head_right}</span></div>'
            f'<div class="entry-sub"><span>{sub_left}</span><span>{sub_right}</span></div>'
            f'<ul>{bullets}</ul>'
            f"</div>"
        )
    blocks.append("</section>")
    return "".join(blocks)


def _build_pdf(parsed: ResumeDocument) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
        title=parsed.name,
    )

    styles = _build_styles()
    story = _build_story(parsed, styles)
    doc.build(story)
    return buffer.getvalue()


def _build_docx(parsed: ResumeDocument) -> bytes:
    document = Document()
    for section in document.sections:
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(parsed.name)
    run.bold = True
    run.font.size = Pt(21)

    meta_parts: list[str] = []
    if parsed.location:
        meta_parts.append(parsed.location)
    if parsed.phone:
        meta_parts.append(parsed.phone)
    if parsed.email:
        meta_parts.append(parsed.email)
    meta_parts.extend(label for label, _ in parsed.links[:5])
    if meta_parts:
        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta.add_run(" | ".join(meta_parts))
        meta_run.font.size = Pt(9.5)

    _docx_add_summary(document, parsed.summary)
    _docx_add_entries_section(document, "EDUCATION", parsed.education)
    _docx_add_entries_section(document, "EXPERIENCE", parsed.experience)
    _docx_add_entries_section(document, "PROJECTS", parsed.projects, project_mode=True)
    _docx_add_skills_section(document, parsed.skills)
    _docx_add_list_section(document, "CERTIFICATIONS", parsed.certifications)
    _docx_add_list_section(document, "EXTRACURRICULAR", parsed.extracurricular)

    out = BytesIO()
    document.save(out)
    return out.getvalue()


def _build_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "ResumeName",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=24,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#0f172a"),
            spaceAfter=4,
        ),
        "meta": ParagraphStyle(
            "ResumeMeta",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=12,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#475569"),
            spaceAfter=8,
        ),
        "section": ParagraphStyle(
            "ResumeSection",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            alignment=TA_LEFT,
            textColor=colors.HexColor("#0f766e"),
            spaceBefore=8,
            spaceAfter=4,
        ),
        "entry_title": ParagraphStyle(
            "EntryTitle",
            parent=base["BodyText"],
            fontName="Helvetica-Bold",
            fontSize=10.2,
            leading=12,
            textColor=colors.HexColor("#0f172a"),
        ),
        "entry_meta": ParagraphStyle(
            "EntryMeta",
            parent=base["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=9.2,
            leading=11,
            textColor=colors.HexColor("#334155"),
        ),
        "body": ParagraphStyle(
            "ResumeBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.2,
            leading=11.4,
            textColor=colors.HexColor("#1e293b"),
            spaceAfter=2,
        ),
        "bullet": ParagraphStyle(
            "ResumeBullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.1,
            leading=11.2,
            leftIndent=10,
            firstLineIndent=-7,
            bulletIndent=0,
            textColor=colors.HexColor("#1e293b"),
            spaceAfter=1.5,
        ),
        "skill": ParagraphStyle(
            "ResumeSkill",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.1,
            leading=11.2,
            textColor=colors.HexColor("#1e293b"),
            spaceAfter=2,
        ),
    }


def _build_story(parsed: ResumeDocument, styles: dict[str, ParagraphStyle]) -> list:
    story: list = []

    story.append(Paragraph(_pdf_escape(parsed.name or "Candidate Name"), styles["name"]))
    meta_parts: list[str] = []
    if parsed.location:
        meta_parts.append(_pdf_escape(parsed.location))
    if parsed.phone:
        meta_parts.append(_pdf_escape(parsed.phone))
    if parsed.email:
        meta_parts.append(_pdf_escape(parsed.email))
    for label, url in parsed.links[:5]:
        meta_parts.append(f'<link href="{escape(url)}">{_pdf_escape(label)}</link>')
    if meta_parts:
        story.append(Paragraph(" | ".join(meta_parts), styles["meta"]))
    story.append(Spacer(1, 4))

    _add_summary(story, parsed.summary, styles)
    _add_entries_section(story, "EDUCATION", parsed.education, styles)
    _add_entries_section(story, "EXPERIENCE", parsed.experience, styles)
    _add_projects_section(story, parsed.projects, styles)
    _add_skills_section(story, parsed.skills, styles)
    _add_list_section(story, "CERTIFICATIONS", parsed.certifications, styles)
    _add_list_section(story, "EXTRACURRICULAR", parsed.extracurricular, styles)

    return story


def _docx_add_section_heading(document: Document, title: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(1)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(3)
    border_run = rule.add_run("_" * 140)
    border_run.font.size = Pt(4)
    border_run.font.color.rgb = RGBColor(203, 213, 225)


def _docx_add_summary(document: Document, summary: list[str]) -> None:
    if not summary:
        return
    _docx_add_section_heading(document, "SUMMARY")
    paragraph = document.add_paragraph(" ".join(summary))
    paragraph.paragraph_format.space_after = Pt(4)


def _docx_add_entries_section(
    document: Document,
    title: str,
    entries: list[ResumeEntry],
    project_mode: bool = False,
) -> None:
    if not entries:
        return
    _docx_add_section_heading(document, title)
    for entry in entries:
        _docx_add_entry(document, entry, project_mode=project_mode)


def _docx_add_entry(document: Document, entry: ResumeEntry, project_mode: bool = False) -> None:
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(5.9)
    table.columns[1].width = Inches(1.1)

    left_title = entry.title
    if project_mode and entry.subtitle:
        left_title = f"{left_title} | {entry.subtitle}"

    top_left = table.cell(0, 0).paragraphs[0]
    top_left.paragraph_format.space_after = Pt(0)
    top_left_run = top_left.add_run(left_title)
    top_left_run.bold = True

    top_right = table.cell(0, 1).paragraphs[0]
    top_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    top_right.paragraph_format.space_after = Pt(0)
    top_right_run = top_right.add_run(entry.date)
    top_right_run.bold = True

    if not project_mode:
        bottom_left = table.cell(1, 0).paragraphs[0]
        bottom_left.paragraph_format.space_after = Pt(0)
        bottom_left_run = bottom_left.add_run(entry.subtitle)
        bottom_left_run.italic = True
    else:
        table.cell(1, 0).text = ""

    bottom_right = table.cell(1, 1).paragraphs[0]
    bottom_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    bottom_right.paragraph_format.space_after = Pt(0)
    bottom_right_run = bottom_right.add_run("" if project_mode else entry.location)
    bottom_right_run.italic = True

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

    for bullet in entry.bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(0.5)
        paragraph.add_run(bullet)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def _docx_add_skills_section(document: Document, skills: list[tuple[str, str]]) -> None:
    if not skills:
        return
    _docx_add_section_heading(document, "TECHNICAL SKILLS")
    for label, value in skills:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        paragraph.add_run(value)


def _docx_add_list_section(document: Document, title: str, items: list[str]) -> None:
    if not items:
        return
    _docx_add_section_heading(document, title)
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(0.5)
        paragraph.add_run(item)


def _add_section_heading(story: list, title: str, styles: dict[str, ParagraphStyle]) -> None:
    story.append(Paragraph(title, styles["section"]))
    story.append(HRFlowable(width="100%", thickness=0.7, color=colors.HexColor("#cbd5e1"), spaceAfter=4))


def _add_summary(story: list, summary: list[str], styles: dict[str, ParagraphStyle]) -> None:
    if not summary:
        return
    _add_section_heading(story, "SUMMARY", styles)
    story.append(Paragraph(_pdf_escape(" ".join(summary)), styles["body"]))


def _add_entries_section(story: list, title: str, entries: list[ResumeEntry], styles: dict[str, ParagraphStyle]) -> None:
    if not entries:
        return
    _add_section_heading(story, title, styles)
    for entry in entries:
        _add_entry(story, entry, styles)


def _add_projects_section(story: list, entries: list[ResumeEntry], styles: dict[str, ParagraphStyle]) -> None:
    if not entries:
        return
    _add_section_heading(story, "PROJECTS", styles)
    for entry in entries:
        _add_entry(story, entry, styles, project_mode=True)


def _add_entry(story: list, entry: ResumeEntry, styles: dict[str, ParagraphStyle], project_mode: bool = False) -> None:
    title_text = _pdf_escape(entry.title)
    if project_mode and entry.subtitle:
        title_text = f"{title_text} <font color='#0f766e'>|</font> {_pdf_escape(entry.subtitle)}"

    head_table = Table(
        [[Paragraph(title_text, styles["entry_title"]), Paragraph(_pdf_escape(entry.date), styles["entry_title"])]],
        colWidths=[4.9 * inch, 1.4 * inch],
        hAlign="LEFT",
    )
    head_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(head_table)

    if entry.subtitle or entry.location:
        meta_table = Table(
            [[Paragraph(_pdf_escape(entry.subtitle), styles["entry_meta"]), Paragraph(_pdf_escape(entry.location), styles["entry_meta"])]],
            colWidths=[4.9 * inch, 1.4 * inch],
            hAlign="LEFT",
        )
        meta_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                ]
            )
        )
        story.append(meta_table)

    for bullet in entry.bullets:
        story.append(Paragraph(_pdf_escape(bullet), styles["bullet"], bulletText="•"))
    story.append(Spacer(1, 4))


def _add_skills_section(story: list, skills: list[tuple[str, str]], styles: dict[str, ParagraphStyle]) -> None:
    if not skills:
        return
    _add_section_heading(story, "TECHNICAL SKILLS", styles)
    for label, value in skills:
        story.append(Paragraph(f"<b>{_pdf_escape(label)}:</b> {_pdf_escape(value)}", styles["skill"]))


def _add_list_section(story: list, title: str, items: list[str], styles: dict[str, ParagraphStyle]) -> None:
    if not items:
        return
    _add_section_heading(story, title, styles)
    for item in items:
        story.append(Paragraph(_pdf_escape(item), styles["bullet"], bulletText="•"))


def _pdf_escape(value: str) -> str:
    return escape(value or "").replace("\n", "<br/>")


def _split_resume_sections(lines: list[str]) -> tuple[list[str], dict[str, list[str]]]:
    header_lines: list[str] = []
    sections: dict[str, list[str]] = {key: [] for key in SECTION_ORDER}
    current_section: str | None = None
    body_lines: list[str] = []
    header_complete = False

    for raw_line in lines:
        line = raw_line.strip()
        section_match = _detect_section_heading(line)
        if section_match:
            current_section, remainder = section_match
            header_complete = True
            if remainder:
                sections[current_section].append(remainder)
            continue

        if current_section is None:
            if not line:
                if header_lines:
                    header_complete = True
                continue
            if not header_complete:
                header_lines.append(line)
            else:
                body_lines.append(raw_line)
            continue

        sections[current_section].append(raw_line)

    if not any(sections[key] for key in SECTION_ORDER) and body_lines:
        inferred = _infer_sections_from_body(body_lines)
        for key in SECTION_ORDER:
            sections[key] = inferred.get(key, [])

    return header_lines, sections


def _normalize_heading(value: str) -> str:
    cleaned = _strip_markdown_artifacts(value).strip()
    cleaned = re.sub(r"^[#*\-\s`_>]+", "", cleaned)
    cleaned = re.sub(r"[#*`_]+$", "", cleaned)
    cleaned = cleaned.strip().strip(":").strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.lower()


def _detect_section_heading(value: str) -> tuple[str, str] | None:
    normalized = _normalize_heading(value)
    if not normalized:
        return None

    if normalized in SECTION_ALIASES:
        return SECTION_ALIASES[normalized], ""

    for alias, canonical in SECTION_ALIASES.items():
        if normalized.startswith(f"{alias}:"):
            remainder = normalized[len(alias) + 1 :].strip()
            return canonical, remainder

    return None


def _parse_header(lines: list[str]) -> dict[str, object]:
    non_empty = [line.strip() for line in lines if line.strip()]
    name = non_empty[0] if non_empty else "Candidate Name"
    remainder = non_empty[1:]

    joined = " | ".join(remainder)
    email_match = EMAIL_RE.search(joined)
    phone_match = PHONE_RE.search(joined)
    urls = URL_RE.findall(joined)

    email = email_match.group(0) if email_match else ""
    phone = phone_match.group(0).strip() if phone_match else ""

    links: list[tuple[str, str]] = []
    for url in urls:
        normalized_url = url if url.startswith("http") else f"https://{url}"
        label = _label_for_url(normalized_url)
        if (label, normalized_url) not in links:
            links.append((label, normalized_url))

    location = ""
    for line in remainder:
        probe = URL_RE.sub("", EMAIL_RE.sub("", PHONE_RE.sub("", line))).strip(" |-")
        if probe:
            location = probe
            break

    return {
        "name": name,
        "location": location,
        "phone": phone,
        "email": email,
        "links": links,
    }


def _label_for_url(url: str) -> str:
    lowered = url.lower()
    if "linkedin.com" in lowered:
        return "LinkedIn"
    if "github.com" in lowered:
        return "GitHub"
    if "leetcode.com" in lowered:
        return "LeetCode"
    if "codeforces.com" in lowered:
        return "Codeforces"
    if "geeksforgeeks.org" in lowered:
        return "GeeksforGeeks"
    domain = lowered.split("//", 1)[-1].split("/", 1)[0]
    return domain.replace("www.", "")


def _parse_simple_lines(lines: list[str]) -> list[str]:
    values: list[str] = []
    for line in lines:
        cleaned = _clean_bullet(line)
        if cleaned:
            values.append(cleaned)
    return values


def _infer_sections_from_body(lines: list[str]) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {key: [] for key in SECTION_ORDER}
    blocks = _group_into_blocks(lines)

    for index, block in enumerate(blocks):
        if not block:
            continue
        section = _classify_block(block, first_body_block=index == 0)
        if section:
            sections[section].extend(block)

    return sections


def _classify_block(block: list[str], *, first_body_block: bool) -> str:
    text = "\n".join(block)
    lowered = text.lower()
    bullet_count = sum(1 for line in block if _is_bullet_line(line))
    has_dates = any(DATE_RE.search(line) for line in block)

    if _looks_like_skills_block(block):
        return "skills"
    if any(keyword in lowered for keyword in ("b.tech", "btech", "m.tech", "bachelor", "master", "university", "college", "cgpa", "gpa")):
        return "education"
    if any(keyword in lowered for keyword in ("certification", "certifications", "certified", "udemy", "coursera")):
        return "certifications"
    if any(keyword in lowered for keyword in ("project", "github.com", "built ", "developed ")) and not has_dates:
        if bullet_count > 0 or "|" in block[0]:
            return "projects"
    if has_dates or bullet_count > 0:
        return "experience"
    if first_body_block:
        return "summary"
    return "extracurricular"


def _looks_like_skills_block(block: list[str]) -> bool:
    short_colon_lines = 0
    for line in block:
        cleaned = _clean_bullet(line)
        if ":" not in cleaned:
            continue
        label, value = cleaned.split(":", 1)
        if label.strip() and value.strip() and len(label.strip()) <= 30:
            short_colon_lines += 1
    return short_colon_lines >= 1


def _parse_skills(lines: list[str]) -> list[tuple[str, str]]:
    skills: list[tuple[str, str]] = []
    uncategorized: list[str] = []

    for line in lines:
        cleaned = _clean_bullet(line)
        if not cleaned:
            continue
        if ":" in cleaned:
            label, value = cleaned.split(":", 1)
            label = label.strip()
            value = value.strip()
            if label and value:
                skills.append((label, value))
                continue
        uncategorized.append(cleaned)

    if uncategorized:
        skills.append(("Core Skills", ", ".join(uncategorized)))
    return skills


def _parse_standard_entries(lines: list[str], *, kind: str) -> list[ResumeEntry]:
    blocks = _group_into_blocks(lines)
    entries = [_parse_entry_block(block, kind=kind) for block in blocks]
    return [entry for entry in entries if entry.title or entry.subtitle or entry.bullets]


def _group_into_blocks(lines: list[str]) -> list[list[str]]:
    blocks: list[list[str]] = []
    current: list[str] = []

    for raw_line in lines:
        stripped = raw_line.strip()
        if not stripped:
            if current:
                blocks.append(current)
                current = []
            continue
        if current and not _is_bullet_line(stripped) and any(_is_bullet_line(line) for line in current):
            blocks.append(current)
            current = [stripped]
            continue
        current.append(stripped)

    if current:
        blocks.append(current)
    return blocks


def _parse_entry_block(block: list[str], *, kind: str) -> ResumeEntry:
    header_lines = [line for line in block if not _is_bullet_line(line)]
    bullet_lines = [_clean_bullet(line) for line in block if _is_bullet_line(line)]
    bullet_lines = [line for line in bullet_lines if line]

    title = ""
    subtitle = ""
    date = ""
    location = ""

    if header_lines:
        title, date = _split_date(header_lines[0])
        if kind != "projects":
            title, subtitle = _split_title_subtitle(title)
        if len(header_lines) > 1:
            secondary_subtitle, secondary_location = _parse_secondary_header(header_lines[1], kind=kind)
            if secondary_subtitle and not subtitle:
                subtitle = secondary_subtitle
            location = secondary_location

    if kind == "projects" and header_lines:
        title, subtitle = _split_project_title_and_stack(title)
        if len(header_lines) > 1 and not date:
            _, date = _split_date(header_lines[1])

    return ResumeEntry(
        title=title.strip(),
        date=date.strip(),
        subtitle=subtitle.strip(),
        location=location.strip(),
        bullets=bullet_lines,
    )


def _split_date(text: str) -> tuple[str, str]:
    match = DATE_RE.search(text)
    if not match:
        return text.strip(), ""
    date = match.group(0).strip(" -|")
    start, end = match.span()
    before = text[:start].rstrip(" -|,")
    after = text[end:].lstrip(" -|,")
    combined = before if not after else f"{before} | {after}" if before else after
    return combined.strip(), date


def _split_location(text: str) -> tuple[str, str]:
    parts = [part.strip() for part in re.split(r"\s+\|\s+|,\s+(?=[A-Z])", text) if part.strip()]
    if len(parts) >= 2:
        return parts[0], ", ".join(parts[1:])
    return text.strip(), ""


def _parse_secondary_header(text: str, *, kind: str) -> tuple[str, str]:
    cleaned = text.strip()
    if kind in {"experience", "education"} and "|" not in cleaned and " at " not in cleaned.lower():
        return "", cleaned
    return _split_location(cleaned)


def _split_title_subtitle(text: str) -> tuple[str, str]:
    if " at " in text.lower():
        parts = re.split(r"(?i)\s+at\s+", text, maxsplit=1)
        if len(parts) == 2:
            return parts[1].strip(), parts[0].strip()
    if "|" in text:
        parts = [part.strip() for part in text.split("|") if part.strip()]
        if len(parts) >= 2:
            return parts[0], " | ".join(parts[1:])
    if " - " in text:
        parts = [part.strip() for part in text.split(" - ", 1)]
        if len(parts) == 2:
            return parts[0], parts[1]
    return text, ""


def _split_project_title_and_stack(text: str) -> tuple[str, str]:
    if "|" in text:
        parts = [part.strip() for part in text.split("|") if part.strip()]
        if len(parts) >= 2:
            return parts[0], " | ".join(parts[1:])
    return text, ""


def _is_bullet_line(line: str) -> bool:
    return bool(re.match(r"^[-*•]\s+", line.strip()))


def _clean_bullet(line: str) -> str:
    stripped = _strip_markdown_artifacts(line).strip()
    return re.sub(r"^[-*•]\s+", "", stripped).strip()


def _strip_markdown_artifacts(text: str) -> str:
    cleaned = text.replace("\u00a0", " ")
    cleaned = re.sub(r"\*\*(.*?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__(.*?)__", r"\1", cleaned)
    cleaned = re.sub(r"`([^`]*)`", r"\1", cleaned)
    cleaned = re.sub(r"^\s*#{1,6}\s*", "", cleaned)
    cleaned = re.sub(r"^\s*>\s*", "", cleaned)
    cleaned = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", cleaned)
    return cleaned
